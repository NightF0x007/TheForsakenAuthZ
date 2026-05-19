#!/usr/bin/env python3
"""
Generate prevalence artifacts from an Incident_Extraction workbook.

Outputs:
- {version}_Prevalence_Snapshot.docx
- {version}_Misconfiguration_Matrix.docx
- {version}_Prevalence_Brief.docx
- {version}_statistics.json
- {version}_validation_report.md
- {version}_sanitized_export.csv (optional via --export_csv)

Usage:
  python generate_artifacts.py --xlsx Incident_Extraction_Codebook_v0.6.1.xlsx \
      --out_dir artifacts --version v0.6.1 --tz America/Los_Angeles --export_csv

  python generate_artifacts.py --xlsx Incident_Extraction_Codebook_v0.6.1.xlsx \
      --out_dir artifacts --version v0.6.1 --validate_only

Design:
- Does not hardcode Codebook v0.6.1 category names.
- Discovers picklists from workbook sheets and list data validations when available.
- Falls back to observed dataset categories when picklists are unavailable.
- Primary-only analysis uses Misconfig_1 / Controls_1.
- Any-occurrence analysis uses Misconfig_1 OR Misconfig_2 and Controls_1 OR Controls_2,
  deduplicated within each incident.
"""
import argparse
import datetime as dt
import json
import logging
import re
from collections import Counter
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from zoneinfo import ZoneInfo

import openpyxl
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl.utils import column_index_from_string, get_column_letter, range_boundaries

COLOR_PRIMARY_BLUE = RGBColor(0x1F, 0x4E, 0x79)
COLOR_SECONDARY_GRAY = RGBColor(0x55, 0x55, 0x55)
COLOR_HEADER_BG = "EDEDED"

FONT_SIZE_TITLE = 24
FONT_SIZE_SUBTITLE = 14
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 13
FONT_SIZE_H3 = 11
FONT_SIZE_NORMAL = 11
FONT_SIZE_SMALL = 10

STANDARD_COL_WIDTHS = [Inches(4.5), Inches(0.75), Inches(0.75)]
MATRIX_COL_WIDTHS = [Inches(1.65), Inches(0.95), Inches(0.95), Inches(1.55), Inches(1.4)]
DEFAULT_CELL_MARGINS = {"top": 80, "start": 100, "bottom": 80, "end": 100}

REQUIRED_COLUMNS = [
    "Incident_ID", "Source_Date", "Attack_Type", "IdP_Context",
    "Impact_Primary", "Entry_Vector", "Source_Org", "Misconfig_1",
    "Misconfig_2", "Controls_1", "Controls_2", "Confidence",
]

OPTIONAL_COLUMNS = [
    "Source_URL", "SaaS_Context", "OAuth_Flow", "Token_Artifacts",
    "Detection_Signals", "Response_Actions", "Notes", "Source_Type",
]

SANITIZED_CSV_COLUMNS = [
    "Incident_ID", "Source_Date", "Year", "Attack_Type", "IdP_Context", "SaaS_Context",
    "Impact_Primary", "Entry_Vector", "Source_Org", "OAuth_Flow", "Token_Artifacts",
    "Misconfig_1", "Misconfig_2", "Controls_1", "Controls_2", "Confidence", "Source_URL",
]

CATEGORY_COLUMNS = [
    "Attack_Type", "Entry_Vector", "IdP_Context", "SaaS_Context", "OAuth_Flow",
    "Impact_Primary", "Misconfig_1", "Misconfig_2", "Controls_1", "Controls_2",
    "Confidence", "Source_Type",
]

PICKLIST_ALIASES = {
    "Attack_Type": ["attacktype", "attack", "oauthabusemechanism"],
    "Entry_Vector": ["entryvector", "entry", "initialaccess"],
    "Misconfig_1": ["misconfig", "misconfiguration", "misconfigurations"],
    "Misconfig_2": ["misconfig", "misconfiguration", "misconfigurations"],
    "Controls_1": ["controls", "control", "controlgap", "controlgaps"],
    "Controls_2": ["controls", "control", "controlgap", "controlgaps"],
    "IdP_Context": ["idpcontext", "idp", "identityprovider"],
    "SaaS_Context": ["saascontext", "saas", "application", "platform"],
    "OAuth_Flow": ["oauthflow", "flow", "granttype"],
    "Impact_Primary": ["impactprimary", "impact", "primaryimpact"],
    "Confidence": ["confidence", "evidencestrength"],
    "Source_Type": ["sourcetype", "source"],
}

logger = logging.getLogger(__name__)


def setup_logging(level: int = logging.INFO) -> logging.Logger:
    logger.setLevel(level)
    if not logger.handlers:
        handler = logging.StreamHandler()
        handler.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
        logger.addHandler(handler)
    return logger


def is_blank(value: Any) -> bool:
    if value is None:
        return True
    try:
        if pd.isna(value):
            return True
    except Exception:
        pass
    return isinstance(value, str) and value.strip() == ""


def clean_str(value: Any) -> str:
    return "" if is_blank(value) else str(value).strip()


def normalize_label(label: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", clean_str(label).lower())


def pct(n: int, total: int) -> float:
    return round((100 * n / total), 1) if total else 0.0


def parse_year(value: Any) -> Optional[int]:
    if is_blank(value):
        return None
    if isinstance(value, (dt.datetime, dt.date)):
        return int(value.year)
    match = re.search(r"(\d{4})", str(value))
    return int(match.group(1)) if match else None


def get_years(series: pd.Series) -> pd.Series:
    years = pd.to_numeric(series.apply(parse_year), errors="coerce").dropna()
    return years.astype(int)


def format_date_range(df: pd.DataFrame) -> str:
    years = get_years(df["Source_Date"])
    if years.empty:
        return "unknown"
    min_year, max_year = int(years.min()), int(years.max())
    return str(min_year) if min_year == max_year else f"{min_year}-{max_year}"


def timestamp_for_timezone(tz_name: str) -> str:
    try:
        tz = ZoneInfo(tz_name)
    except Exception as exc:
        raise ValueError(f"Invalid timezone '{tz_name}'. Use an IANA timezone such as UTC or America/Los_Angeles.") from exc
    return dt.datetime.now(tz).strftime("%Y-%m-%d %H:%M %z")


def json_safe(value: Any) -> Any:
    if isinstance(value, Counter):
        return {str(k): json_safe(v) for k, v in value.items()}
    if isinstance(value, dict):
        return {str(k): json_safe(v) for k, v in value.items()}
    if isinstance(value, (list, tuple, set)):
        return [json_safe(v) for v in value]
    if isinstance(value, (dt.datetime, dt.date)):
        return value.isoformat()
    if value is pd.NA:
        return None
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass
    if hasattr(value, "item"):
        try:
            return value.item()
        except Exception:
            pass
    return value


def nonblank_unique(values: Iterable[Any]) -> List[str]:
    return sorted({clean_str(v) for v in values if not is_blank(v)})


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def apply_doc_style(doc: Document) -> None:
    styles = doc.styles
    if "Title" in styles:
        st = styles["Title"]
        st.font.name = "Calibri"
        st.font.size = Pt(FONT_SIZE_TITLE)
        st.font.bold = True
        st.font.color.rgb = COLOR_PRIMARY_BLUE

    for name, size in {"Heading 1": FONT_SIZE_H1, "Heading 2": FONT_SIZE_H2, "Heading 3": FONT_SIZE_H3}.items():
        if name in styles:
            st = styles[name]
            st.font.name = "Calibri"
            st.font.bold = True
            st.font.color.rgb = COLOR_PRIMARY_BLUE
            st.font.size = Pt(size)

    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(FONT_SIZE_NORMAL)


def add_header_section(doc: Document, title: str, subtitle: str, gen_date: str, meta_text: str) -> None:
    doc.add_paragraph(title, style="Title")
    sub = doc.add_paragraph(subtitle)
    if sub.runs:
        sub.runs[0].font.size = Pt(FONT_SIZE_SUBTITLE)
        sub.runs[0].font.color.rgb = COLOR_PRIMARY_BLUE
    meta = doc.add_paragraph()
    run = meta.add_run(f"Generated: {gen_date}\n{meta_text}")
    run.font.size = Pt(FONT_SIZE_SMALL)
    run.font.color.rgb = COLOR_SECONDARY_GRAY


def add_table(doc: Document, title: Optional[str], rows: List[Tuple[Any, ...]], col_widths: List[Any], header: Tuple[str, ...]) -> Any:
    if title:
        doc.add_paragraph(title, style="Heading 2")

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    style_table(table)

    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = str(h)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
        set_cell_shading(cell, COLOR_HEADER_BG)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = f"{value:.1f}%" if j == 2 and isinstance(value, (float, int)) else str(value)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

    for j, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[j].width = width

    return table


def load_workbook_safe(xlsx_path: str) -> openpyxl.Workbook:
    try:
        logger.info("Loading workbook: %s", xlsx_path)
        return openpyxl.load_workbook(xlsx_path, data_only=True)
    except FileNotFoundError as exc:
        raise FileNotFoundError(f"Workbook not found: {xlsx_path}") from exc
    except Exception as exc:
        raise ValueError(f"Invalid Excel file: {xlsx_path}") from exc


def load_incidents(wb: openpyxl.Workbook) -> pd.DataFrame:
    if "Incident_Extraction" not in wb.sheetnames:
        raise KeyError("Incident_Extraction sheet not found")

    ws = wb["Incident_Extraction"]
    headers = [clean_str(c.value) for c in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row and any(not is_blank(v) for v in row):
            rows.append(row[:len(headers)])

    df = pd.DataFrame(rows, columns=headers)
    missing = sorted(set(REQUIRED_COLUMNS) - set(df.columns))
    if missing:
        raise ValueError(f"Missing required columns: {missing}")

    logger.info("Loaded %s incidents", len(df))
    return df


def read_column_picklists(ws) -> Dict[str, List[str]]:
    picklists: Dict[str, List[str]] = {}
    if ws.max_row < 2:
        return picklists

    headers = [clean_str(c.value) for c in ws[1]]
    for col_idx, header in enumerate(headers, start=1):
        if not header:
            continue
        values = []
        for row_idx in range(2, min(ws.max_row, 1000) + 1):
            value = ws.cell(row_idx, col_idx).value
            if is_blank(value):
                if values:
                    break
                continue
            values.append(clean_str(value))
        if values:
            picklists[header] = values
    return picklists


def parse_range_reference(formula: str) -> Optional[Tuple[str, str]]:
    if not formula:
        return None
    formula = formula.strip().strip('"').lstrip("=")
    match = re.match(r"'?([^'!]+)'?!\$?([A-Z]+)\$?\d+:\$?[A-Z]+\$?\d+", formula)
    if not match:
        return None
    return match.group(1), formula.split("!", 1)[1] if "!" in formula else formula


def values_from_range_reference(wb: openpyxl.Workbook, formula: str) -> List[str]:
    formula = formula.strip().strip('"').lstrip("=")
    if "," in formula and "!" not in formula:
        return [v.strip() for v in formula.split(",") if v.strip()]

    if "!" not in formula:
        return []

    sheet_name, cell_range = formula.split("!", 1)
    sheet_name = sheet_name.strip("'")
    if sheet_name not in wb.sheetnames:
        return []
    ws = wb[sheet_name]
    try:
        min_col, min_row, max_col, max_row = range_boundaries(cell_range.replace("$", ""))
    except Exception:
        return []

    values = []
    for row in ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col, values_only=True):
        for value in row:
            if not is_blank(value):
                values.append(clean_str(value))
    return values


def read_data_validation_picklists(wb: openpyxl.Workbook) -> Dict[str, List[str]]:
    if "Incident_Extraction" not in wb.sheetnames:
        return {}

    ws = wb["Incident_Extraction"]
    headers_by_col = {idx: clean_str(cell.value) for idx, cell in enumerate(ws[1], start=1) if clean_str(cell.value)}
    picklists: Dict[str, List[str]] = {}

    for dv in ws.data_validations.dataValidation:
        if dv.type != "list" or not dv.formula1:
            continue
        values = values_from_range_reference(wb, dv.formula1)
        if not values:
            continue
        for cell_range in dv.sqref.ranges:
            min_col, _, max_col, _ = range_boundaries(str(cell_range))
            for col_idx in range(min_col, max_col + 1):
                header = headers_by_col.get(col_idx)
                if header:
                    picklists[header] = values
    return picklists


def read_picklists_from_workbook(wb: openpyxl.Workbook) -> Dict[str, List[str]]:
    picklists: Dict[str, List[str]] = {}

    for sheet_name in ["Picklists", "Codebook", "Lookup", "Reference"]:
        if sheet_name in wb.sheetnames:
            logger.info("Reading picklists from sheet: %s", sheet_name)
            sheet_picklists = read_column_picklists(wb[sheet_name])
            picklists.update(sheet_picklists)
            break

    validation_picklists = read_data_validation_picklists(wb)
    for key, values in validation_picklists.items():
        picklists.setdefault(key, values)

    if picklists:
        logger.info("Discovered %s picklist group(s)", len(picklists))
    else:
        logger.warning("No picklists discovered; falling back to observed categories")

    return picklists


def find_picklist_values(picklists: Dict[str, List[str]], column_name: str) -> Optional[List[str]]:
    if not picklists:
        return None

    normalized = {normalize_label(k): v for k, v in picklists.items()}
    candidates = [normalize_label(column_name)] + PICKLIST_ALIASES.get(column_name, [])

    for candidate in candidates:
        if candidate in normalized:
            return normalized[candidate]

    for key, values in normalized.items():
        if any(candidate and (candidate in key or key in candidate) for candidate in candidates):
            return values

    return None


class ValidationReport:
    def __init__(self) -> None:
        self.warnings: List[str] = []
        self.errors: List[str] = []
        self.stats: Dict[str, Any] = {}

    def add_warning(self, message: str) -> None:
        self.warnings.append(message)
        logger.warning(message)

    def add_error(self, message: str) -> None:
        self.errors.append(message)
        logger.error(message)

    def to_markdown(self) -> str:
        lines = ["# Data Validation Report\n\n"]
        if self.stats:
            lines.append("## Statistics\n\n")
            for key, value in self.stats.items():
                lines.append(f"- {key}: {value}\n")
        if self.warnings:
            lines.append(f"\n## Warnings ({len(self.warnings)})\n\n")
            for warning in self.warnings:
                lines.append(f"- {warning}\n")
        if self.errors:
            lines.append(f"\n## Errors ({len(self.errors)})\n\n")
            for error in self.errors:
                lines.append(f"- {error}\n")
        if not self.warnings and not self.errors:
            lines.append("\n✓ No issues detected.\n")
        return "".join(lines)


def validate_incidents(df: pd.DataFrame, picklists: Dict[str, List[str]]) -> ValidationReport:
    report = ValidationReport()
    report.stats["Total incidents"] = int(len(df))

    duplicate_count = int(df["Incident_ID"].duplicated().sum())
    if duplicate_count:
        report.add_error(f"{duplicate_count} duplicate Incident_ID values found")

    for column in ["Misconfig_1", "Controls_1"]:
        blank_count = int(df[column].apply(is_blank).sum())
        if blank_count:
            report.add_warning(f"{blank_count} blank values in {column}")

    parsed_dates = pd.to_datetime(df["Source_Date"], errors="coerce")
    invalid_dates = int((df["Source_Date"].notna() & parsed_dates.isna()).sum())
    if invalid_dates:
        report.add_warning(f"{invalid_dates} rows with unparseable Source_Date")

    for c1, c2 in [("Misconfig_1", "Misconfig_2"), ("Controls_1", "Controls_2")]:
        bad = df[c1].apply(is_blank) & ~df[c2].apply(is_blank)
        if int(bad.sum()):
            report.add_warning(f"{int(bad.sum())} rows with {c2} filled but {c1} blank")

        duplicate_pair = (
            df[c1].apply(clean_str).eq(df[c2].apply(clean_str))
            & ~df[c1].apply(is_blank)
            & ~df[c2].apply(is_blank)
        )
        if int(duplicate_pair.sum()):
            report.add_warning(f"{int(duplicate_pair.sum())} rows where {c1} and {c2} have identical values")

    for column in CATEGORY_COLUMNS:
        if column not in df.columns:
            continue
        allowed = find_picklist_values(picklists, column)
        if not allowed:
            continue
        observed = set(nonblank_unique(df[column]))
        unexpected = sorted(observed - set(allowed))
        if unexpected:
            preview = "; ".join(unexpected[:10])
            suffix = " ..." if len(unexpected) > 10 else ""
            report.add_warning(
                f"{len(unexpected)} observed value(s) in {column} absent from discovered picklist: {preview}{suffix}"
            )

    conf_counts = df["Confidence"].apply(clean_str).replace("", "Uncoded / Blank").value_counts()
    report.stats["Confidence distribution"] = {str(k): int(v) for k, v in conf_counts.items()}

    years = get_years(df["Source_Date"])
    if not years.empty:
        report.stats["Date range (years)"] = f"{int(years.min())}-{int(years.max())}"

    if picklists:
        report.stats["Discovered picklist groups"] = int(len(picklists))
        report.stats["Picklist group names"] = sorted(str(k) for k in picklists.keys())
    else:
        report.add_warning("No picklists discovered; validation used observed data categories only")

    return report


def counter_for_column(df: pd.DataFrame, column: str) -> Counter:
    if column not in df.columns:
        return Counter()
    return Counter(clean_str(value) for value in df[column] if not is_blank(value))


def any_occurrence_counter(df: pd.DataFrame, columns: Sequence[str]) -> Counter:
    counter = Counter()
    for _, row in df.iterrows():
        row_values = {clean_str(row[col]) for col in columns if col in df.columns and clean_str(row[col])}
        for value in row_values:
            counter[value] += 1
    return counter


def compute_counters(df: pd.DataFrame) -> Dict[str, Counter]:
    counters = {
        "attack_type": counter_for_column(df, "Attack_Type"),
        "entry_vector": counter_for_column(df, "Entry_Vector"),
        "idp_context": counter_for_column(df, "IdP_Context"),
        "saas_context": counter_for_column(df, "SaaS_Context"),
        "impact_primary": counter_for_column(df, "Impact_Primary"),
        "source_org": counter_for_column(df, "Source_Org"),
        "confidence": counter_for_column(df, "Confidence"),
        "oauth_flow": counter_for_column(df, "OAuth_Flow"),
        "misconfig_primary": counter_for_column(df, "Misconfig_1"),
        "misconfig_any": any_occurrence_counter(df, ["Misconfig_1", "Misconfig_2"]),
        "controls_primary": counter_for_column(df, "Controls_1"),
        "controls_any": any_occurrence_counter(df, ["Controls_1", "Controls_2"]),
    }
    return counters


def rows_from_counter(
    counter: Counter,
    denominator: int,
    categories: Optional[List[str]] = None,
    include_zero_picklist_values: bool = False,
    sort_by_count: bool = True,
) -> List[Tuple[str, int, float]]:
    """Build prevalence rows from a Counter.

    When picklist categories are supplied, they are used only for tie-breaking
    and optional zero-value inclusion. Prevalence outputs still sort high-to-low
    by count so executive highlights and tables identify the true top values.
    """
    if categories:
        picklist_index = {category: i for i, category in enumerate(categories)}
        category_set = set(categories)
        keys = [c for c in categories if include_zero_picklist_values or counter.get(c, 0) > 0]
        keys += [k for k in counter if k not in category_set]
        if sort_by_count:
            keys = sorted(keys, key=lambda x: (-int(counter.get(x, 0)), picklist_index.get(x, 9999), x))
        else:
            keys = [k for k in categories if include_zero_picklist_values or counter.get(k, 0) > 0]
            keys += sorted([k for k in counter if k not in category_set], key=lambda x: (-int(counter.get(x, 0)), x))
    else:
        keys = sorted(counter.keys(), key=lambda x: (-int(counter[x]), x))
    return [(key, int(counter.get(key, 0)), pct(int(counter.get(key, 0)), denominator)) for key in keys]

def top_rows(rows: List[Tuple[str, int, float]], n: int = 2) -> List[Tuple[str, int, float]]:
    return rows[:n]


def representative_ids_for_misconfig(df: pd.DataFrame, category: str, max_n: int = 4) -> List[str]:
    sub = df[(df["Misconfig_1"].apply(clean_str) == category) | (df["Misconfig_2"].apply(clean_str) == category)].copy()
    if sub.empty:
        return []

    confidence_order = {"High": 0, "Medium": 1, "Low": 2}
    sub["_confidence_rank"] = sub["Confidence"].apply(lambda x: confidence_order.get(clean_str(x), 9))
    sub["_date"] = pd.to_datetime(sub["Source_Date"], errors="coerce")
    sub = sub.sort_values(by=["_confidence_rank", "_date", "Incident_ID"], ascending=[True, False, True])
    return sub["Incident_ID"].astype(str).head(max_n).tolist()


def common_controls_for_misconfig(df: pd.DataFrame, category: str, max_n: int = 3) -> List[str]:
    sub = df[(df["Misconfig_1"].apply(clean_str) == category) | (df["Misconfig_2"].apply(clean_str) == category)]
    counter = Counter()
    for _, row in sub.iterrows():
        controls = {
            clean_str(row[col])
            for col in ["Controls_1", "Controls_2"]
            if col in df.columns and clean_str(row[col])
        }
        for control in controls:
            counter[control] += 1
    return [key for key, _ in sorted(counter.items(), key=lambda item: (-item[1], item[0]))[:max_n]]


def table_rows_for_field(counters: Dict[str, Counter], key: str, denominator: int, picklists: Dict[str, List[str]], column_name: str) -> List[Tuple[str, int, float]]:
    return rows_from_counter(counters[key], denominator, find_picklist_values(picklists, column_name))


def build_prevalence_snapshot(
    doc: Document,
    df: pd.DataFrame,
    counters: Dict[str, Counter],
    picklists: Dict[str, List[str]],
    gen_date: str,
    version: str,
    workbook_name: str,
) -> None:
    total = len(df)
    date_range = format_date_range(df)

    attack_rows = table_rows_for_field(counters, "attack_type", total, picklists, "Attack_Type")
    entry_rows = table_rows_for_field(counters, "entry_vector", total, picklists, "Entry_Vector")
    idp_rows = table_rows_for_field(counters, "idp_context", total, picklists, "IdP_Context")
    saas_rows = table_rows_for_field(counters, "saas_context", total, picklists, "SaaS_Context")
    impact_rows = table_rows_for_field(counters, "impact_primary", total, picklists, "Impact_Primary")
    org_rows = rows_from_counter(counters["source_org"], total)
    confidence_rows = table_rows_for_field(counters, "confidence", total, picklists, "Confidence")
    mis_primary_rows = rows_from_counter(counters["misconfig_primary"], total, find_picklist_values(picklists, "Misconfig_1"))
    mis_any_rows = rows_from_counter(counters["misconfig_any"], total, find_picklist_values(picklists, "Misconfig_1"))
    ctrl_primary_rows = rows_from_counter(counters["controls_primary"], total, find_picklist_values(picklists, "Controls_1"))
    ctrl_any_rows = rows_from_counter(counters["controls_any"], total, find_picklist_values(picklists, "Controls_1"))

    years = get_years(df["Source_Date"])
    year_counts = Counter(years)
    year_rows = [(str(year), int(year_counts[year]), pct(int(year_counts[year]), total)) for year in sorted(year_counts)]

    add_header_section(
        doc,
        f"Prevalence Snapshot ({version})",
        f"Mapping and Mitigating SaaS OAuth Post‑SSO Abuse ({date_range})",
        gen_date,
        f"Data source: {workbook_name}\nAnalysis note: primary-only = _1 fields; any-occurrence = _1 OR _2, deduplicated per incident.",
    )

    doc.add_paragraph("Executive highlights", style="Heading 1")
    highlights = [f"Total incidents coded in this curated public-reporting dataset: {total} ({date_range})."]

    if attack_rows:
        highlights.append("Top attack types: " + "; ".join(f"{k} ({n}/{total}, {p:.1f}%)" for k, n, p in top_rows(attack_rows)) + ".")
    if impact_rows:
        highlights.append("Top primary impacts: " + "; ".join(f"{k} ({n}/{total}, {p:.1f}%)" for k, n, p in top_rows(impact_rows)) + ".")
    if mis_primary_rows:
        k, n, p = mis_primary_rows[0]
        highlights.append(f"Top primary-only misconfiguration: {k} ({n}/{total}, {p:.1f}%).")
    if mis_any_rows:
        k, n, p = mis_any_rows[0]
        highlights.append(f"Top any-occurrence misconfiguration: {k} ({n}/{total}, {p:.1f}%).")
    if ctrl_primary_rows:
        k, n, p = ctrl_primary_rows[0]
        highlights.append(f"Top primary-only control gap: {k} ({n}/{total}, {p:.1f}%).")
    if ctrl_any_rows:
        k, n, p = ctrl_any_rows[0]
        highlights.append(f"Top any-occurrence control gap: {k} ({n}/{total}, {p:.1f}%).")

    for highlight in highlights:
        doc.add_paragraph(highlight, style="List Bullet")

    doc.add_paragraph("Dataset overview", style="Heading 1")
    if year_rows:
        add_table(doc, "By year (Source_Date)", year_rows, STANDARD_COL_WIDTHS, ("Year", "n", "%"))
    if org_rows:
        add_table(doc, "By source organization (Source_Org)", org_rows, STANDARD_COL_WIDTHS, ("Source_Org", "n", "%"))

    doc.add_paragraph("Core prevalence (all incidents)", style="Heading 1")
    for title, rows, header_name in [
        ("Attack_Type", attack_rows, "Attack_Type"),
        ("Entry_Vector", entry_rows, "Entry_Vector"),
        ("IdP_Context", idp_rows, "IdP_Context"),
        ("SaaS_Context", saas_rows, "SaaS_Context"),
        ("Impact_Primary", impact_rows, "Impact_Primary"),
        ("Confidence", confidence_rows, "Confidence"),
    ]:
        if rows:
            add_table(doc, title, rows, STANDARD_COL_WIDTHS, (header_name, "n", "%"))

    doc.add_paragraph("Misconfiguration prevalence", style="Heading 1")
    if mis_primary_rows:
        add_table(doc, "Misconfiguration (primary-only: Misconfig_1)", mis_primary_rows, STANDARD_COL_WIDTHS, ("Misconfiguration", "n", "%"))
    if mis_any_rows:
        add_table(doc, "Misconfiguration (any-occurrence: Misconfig_1 OR Misconfig_2)", mis_any_rows, STANDARD_COL_WIDTHS, ("Misconfiguration", "n", "%"))

    doc.add_paragraph("Controls prevalence", style="Heading 1")
    if ctrl_primary_rows:
        add_table(doc, "Controls (primary-only: Controls_1)", ctrl_primary_rows, STANDARD_COL_WIDTHS, ("Controls", "n", "%"))
    if ctrl_any_rows:
        add_table(doc, "Controls (any-occurrence: Controls_1 OR Controls_2)", ctrl_any_rows, STANDARD_COL_WIDTHS, ("Controls", "n", "%"))

    logger.info("Prevalence Snapshot built")


def build_misconfiguration_matrix(
    doc: Document,
    df: pd.DataFrame,
    counters: Dict[str, Counter],
    picklists: Dict[str, List[str]],
    gen_date: str,
    version: str,
    workbook_name: str,
) -> None:
    total = len(df)
    picklist_order = find_picklist_values(picklists, "Misconfig_1") or []
    observed = set(counters["misconfig_any"].keys())
    categories = [c for c in picklist_order if c in observed]
    categories += sorted(observed - set(categories))

    add_header_section(
        doc,
        f"Misconfiguration Matrix ({version})",
        f"Mapping and Mitigating SaaS OAuth Post‑SSO Abuse ({format_date_range(df)})",
        gen_date,
        f"Data source: {workbook_name}\nPrimary-only uses Misconfig_1. Any-occurrence uses Misconfig_1 OR Misconfig_2, deduplicated per incident.",
    )

    doc.add_paragraph("Misconfiguration categories", style="Heading 1")
    doc.add_paragraph(
        "Category definitions are maintained in the source workbook/codebook. This generated matrix does not invent or embed definitions; it reports prevalence, representative incident identifiers, and observed control co-occurrence."
    )

    headers = ("Misconfiguration", "Primary n (%)", "Any n (%)", "Representative incidents", "Common controls (top)")
    table = doc.add_table(rows=1 + len(categories), cols=len(headers))
    style_table(table)

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
        set_cell_shading(cell, COLOR_HEADER_BG)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT

    for i, category in enumerate(categories, start=1):
        primary_n = int(counters["misconfig_primary"].get(category, 0))
        any_n = int(counters["misconfig_any"].get(category, 0))
        reps = ", ".join(representative_ids_for_misconfig(df, category))
        controls = ", ".join(common_controls_for_misconfig(df, category))

        row = table.rows[i].cells
        row[0].text = category
        row[1].text = f"{primary_n} ({pct(primary_n, total):.1f}%)"
        row[2].text = f"{any_n} ({pct(any_n, total):.1f}%)"
        row[3].text = reps
        row[4].text = controls
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for j, width in enumerate(MATRIX_COL_WIDTHS):
        for row in table.rows:
            row.cells[j].width = width

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(FONT_SIZE_SMALL)

    logger.info("Misconfiguration Matrix built")


def build_prevalence_brief(
    doc: Document,
    df: pd.DataFrame,
    counters: Dict[str, Counter],
    picklists: Dict[str, List[str]],
    gen_date: str,
    version: str,
    workbook_name: str,
) -> None:
    total = len(df)
    date_range = format_date_range(df)

    attack_rows = table_rows_for_field(counters, "attack_type", total, picklists, "Attack_Type")
    impact_rows = table_rows_for_field(counters, "impact_primary", total, picklists, "Impact_Primary")
    mis_primary_rows = rows_from_counter(counters["misconfig_primary"], total, find_picklist_values(picklists, "Misconfig_1"))
    mis_any_rows = rows_from_counter(counters["misconfig_any"], total, find_picklist_values(picklists, "Misconfig_1"))
    ctrl_primary_rows = rows_from_counter(counters["controls_primary"], total, find_picklist_values(picklists, "Controls_1"))
    ctrl_any_rows = rows_from_counter(counters["controls_any"], total, find_picklist_values(picklists, "Controls_1"))

    add_header_section(
        doc,
        f"Prevalence Brief ({version})",
        f"SaaS OAuth Post‑SSO Abuse in Public Reporting ({date_range})",
        gen_date,
        f"Dataset: {total} incidents\nSource: {workbook_name}",
    )

    doc.add_paragraph("Scope and method", style="Heading 1")
    doc.add_paragraph(
        f"This brief summarizes {total} entries in a curated public-reporting dataset ({date_range}) that explicitly describe OAuth-enabled post‑SSO abuse in SaaS/IdP contexts. Each row is treated as one coded incident, campaign, or case study. Prevalence is reported in two views: primary-only, which uses Misconfig_1 and Controls_1, and any-occurrence, which treats a category as present if it appears in either _1 or _2 while deduplicating within each incident."
    )

    doc.add_paragraph("Key findings", style="Heading 1")
    findings = [f"Total incidents coded: {total} ({date_range})."]
    if attack_rows:
        findings.append("Top attack types: " + "; ".join(f"{k} ({n}/{total}, {p:.1f}%)" for k, n, p in top_rows(attack_rows)) + ".")
    if impact_rows:
        findings.append("Top primary impacts: " + "; ".join(f"{k} ({n}/{total}, {p:.1f}%)" for k, n, p in top_rows(impact_rows)) + ".")
    if mis_primary_rows:
        k, n, p = mis_primary_rows[0]
        findings.append(f"Top primary-only misconfiguration: {k} ({n}/{total}, {p:.1f}%).")
    if mis_any_rows:
        k, n, p = mis_any_rows[0]
        findings.append(f"Top any-occurrence misconfiguration: {k} ({n}/{total}, {p:.1f}%).")
    if ctrl_primary_rows:
        k, n, p = ctrl_primary_rows[0]
        findings.append(f"Top primary-only control gap: {k} ({n}/{total}, {p:.1f}%).")
    if ctrl_any_rows:
        k, n, p = ctrl_any_rows[0]
        findings.append(f"Top any-occurrence control gap: {k} ({n}/{total}, {p:.1f}%).")

    for finding in findings:
        doc.add_paragraph(finding, style="List Bullet")

    doc.add_paragraph("Prevalence tables", style="Heading 1")
    for title, rows, header in [
        ("Attack_Type", attack_rows, "Attack_Type"),
        ("Impact_Primary", impact_rows, "Impact_Primary"),
        ("Misconfiguration (primary-only)", mis_primary_rows, "Misconfiguration"),
        ("Misconfiguration (any-occurrence)", mis_any_rows, "Misconfiguration"),
        ("Controls (primary-only)", ctrl_primary_rows, "Controls"),
        ("Controls (any-occurrence)", ctrl_any_rows, "Controls"),
    ]:
        if rows:
            add_table(doc, title, rows, STANDARD_COL_WIDTHS, (header, "n", "%"))

    doc.add_paragraph("Limitations", style="Heading 1")
    limitations = [
        "Public-reporting bias: distributions in this dataset should not be interpreted as true global incident prevalence.",
        "Vendor and platform skew: public reporting can overrepresent ecosystems with more mature or more transparent incident reporting.",
        "Outcome ambiguity: some sources describe attempted or campaign-level activity with limited confirmation of downstream impact.",
        "Taxonomy fit: structured fields simplify complex incident narratives; important caveats should remain preserved in source Notes.",
    ]
    for limitation in limitations:
        doc.add_paragraph(limitation, style="List Bullet")

    logger.info("Prevalence Brief built")


def export_sanitized_csv(df: pd.DataFrame, out_path: Path) -> None:
    export_df = df.copy()
    export_df["Year"] = export_df["Source_Date"].apply(parse_year)
    safe_cols = [col for col in SANITIZED_CSV_COLUMNS if col in export_df.columns]
    export_df[safe_cols].to_csv(out_path, index=False)
    logger.info("Sanitized CSV exported: %s", out_path)


def export_statistics_json(
    df: pd.DataFrame,
    counters: Dict[str, Counter],
    report: ValidationReport,
    picklists: Dict[str, List[str]],
    version: str,
    out_path: Path,
) -> None:
    years = get_years(df["Source_Date"])
    stats = {
        "version": version,
        "timestamp_utc": dt.datetime.now(dt.timezone.utc).isoformat(),
        "total_incidents": int(len(df)),
        "date_range": {
            "min_year": int(years.min()) if not years.empty else None,
            "max_year": int(years.max()) if not years.empty else None,
        },
        "counters": {key: dict(counter) for key, counter in counters.items()},
        "picklists": {
            "discovered_group_count": int(len(picklists)),
            "groups": {key: values for key, values in picklists.items()},
        },
        "validation": {
            "warnings": report.warnings,
            "errors": report.errors,
            "stats": report.stats,
        },
    }

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(json_safe(stats), f, indent=2, ensure_ascii=False)
    logger.info("Statistics JSON exported: %s", out_path)


def write_validation_report(report: ValidationReport, out_path: Path) -> None:
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(report.to_markdown())
    logger.info("Validation report exported: %s", out_path)


def build_docs(
    df: pd.DataFrame,
    picklists: Dict[str, List[str]],
    out_dir: str,
    version: str,
    workbook_name: str,
    timezone: str,
    validate_only: bool,
    export_csv: bool,
) -> ValidationReport:
    out_path = Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    gen_date = timestamp_for_timezone(timezone)
    report = validate_incidents(df, picklists)
    counters = compute_counters(df)

    if not validate_only:
        doc = Document()
        apply_doc_style(doc)
        build_prevalence_snapshot(doc, df, counters, picklists, gen_date, version, workbook_name)
        snapshot_file = out_path / f"{version}_Prevalence_Snapshot.docx"
        doc.save(snapshot_file)
        logger.info("Saved: %s", snapshot_file)

        matrix_doc = Document()
        apply_doc_style(matrix_doc)
        build_misconfiguration_matrix(matrix_doc, df, counters, picklists, gen_date, version, workbook_name)
        matrix_file = out_path / f"{version}_Misconfiguration_Matrix.docx"
        matrix_doc.save(matrix_file)
        logger.info("Saved: %s", matrix_file)

        brief_doc = Document()
        apply_doc_style(brief_doc)
        build_prevalence_brief(brief_doc, df, counters, picklists, gen_date, version, workbook_name)
        brief_file = out_path / f"{version}_Prevalence_Brief.docx"
        brief_doc.save(brief_file)
        logger.info("Saved: %s", brief_file)
    else:
        logger.info("Validate-only mode: skipping DOCX generation")

    if export_csv:
        export_sanitized_csv(df, out_path / f"{version}_sanitized_export.csv")

    export_statistics_json(df, counters, report, picklists, version, out_path / f"{version}_statistics.json")
    write_validation_report(report, out_path / f"{version}_validation_report.md")
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate prevalence artifacts from an Incident_Extraction workbook.")
    parser.add_argument("--xlsx", required=True, help="Path to Incident_Extraction workbook")
    parser.add_argument("--out_dir", default=".", help="Output directory for generated artifacts")
    parser.add_argument("--version", default="v0.6.1", help="Schema/codebook version label")
    parser.add_argument("--log_level", choices=["DEBUG", "INFO", "WARNING", "ERROR"], default="INFO")
    parser.add_argument("--validate_only", action="store_true", help="Validate data and emit JSON/report only")
    parser.add_argument("--export_csv", action="store_true", help="Export sanitized analytical CSV")
    parser.add_argument("--tz", default="UTC", help="IANA timezone for document timestamps")
    args = parser.parse_args()

    setup_logging(getattr(logging, args.log_level))

    try:
        workbook_name = Path(args.xlsx).name
        logger.info("Starting artifact generation for %s", workbook_name)
        wb = load_workbook_safe(args.xlsx)
        picklists = read_picklists_from_workbook(wb)
        df = load_incidents(wb)
        report = build_docs(
            df=df,
            picklists=picklists,
            out_dir=args.out_dir,
            version=args.version,
            workbook_name=workbook_name,
            timezone=args.tz,
            validate_only=args.validate_only,
            export_csv=args.export_csv,
        )
        if report.errors:
            logger.warning("Generation completed with %s validation error(s)", len(report.errors))
        logger.info("Generation complete")
    except Exception as exc:
        logger.error("Fatal error: %s", exc, exc_info=True)
        raise SystemExit(1)


if __name__ == "__main__":
    main()
